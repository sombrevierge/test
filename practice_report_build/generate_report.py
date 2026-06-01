from pathlib import Path
from io import BytesIO
import zipfile
import requests
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('practice_report_build/output')
OUT.mkdir(parents=True, exist_ok=True)
ASSETS = Path('practice_report_build/assets')
ASSETS.mkdir(parents=True, exist_ok=True)
DOCX = OUT / 'report_practice_pm01_trubacheva_ev.docx'
INTERFACE_URL = 'https://help.ascon.ru/KOMPAS/21/ru-RU/images/v19_okno_zoom75.png'
DRAWING_URL = 'https://cdn.vmasshtabe.ru/uploads/2015/11/319232-vms-111.png'


def download(url, path):
    r = requests.get(url, timeout=30, headers={'User-Agent': 'Mozilla/5.0'})
    r.raise_for_status()
    path.write_bytes(r.content)


def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.bold = bold
    set_run_font(r, 12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=14, bold=None, italic=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(p, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.5):
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = line
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)
    for r in p.runs:
        set_run_font(r, 14)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, 14, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, 14)
    else:
        r = p.add_run(text)
        set_run_font(r, 14)
    return style_paragraph(p)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, 14, bold=True)
    return p


def add_center(doc, text, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    set_run_font(r, size, bold=bold)
    return p


def add_right(doc, text, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    set_run_font(r, size, bold=bold)
    return p


def add_page_number(section, start=3):
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    set_run_font(run, 12)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.extend([fld_begin, instr, fld_end])


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    rep = OxmlElement('w:tblHeader')
    rep.set(qn('w:val'), 'true')
    trPr.append(rep)


def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement('w:cantSplit')
    trPr.append(cant)


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Cm(width)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    set_run_font(r, 14)
    return p


def create_flowchart(path):
    w, h = 1700, 620
    img = Image.new('RGB', (w, h), 'white')
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf', 34)
        font_b = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf', 34)
    except Exception:
        font = ImageFont.load_default()
        font_b = font
    boxes = [
        ('1', 'Анализ исходных\nданных'),
        ('2', 'Построение\nэскиза'),
        ('3', 'Создание\n3D-модели'),
        ('4', 'Формирование\nассоциативного чертежа'),
        ('5', 'Нанесение размеров\nи проверка оформления'),
    ]
    x0, y0, bw, bh, gap = 35, 175, 280, 235, 48
    for i, (n, txt) in enumerate(boxes):
        x = x0 + i * (bw + gap)
        d.rounded_rectangle([x, y0, x + bw, y0 + bh], radius=24, outline='black', width=4)
        d.ellipse([x+18, y0+18, x+78, y0+78], outline='black', width=3)
        bb = d.textbbox((0,0), n, font=font_b)
        d.text((x+48-(bb[2]-bb[0])/2, y0+48-(bb[3]-bb[1])/2-4), n, fill='black', font=font_b)
        lines = txt.split('\n')
        yy = y0 + 112
        for line in lines:
            bb = d.textbbox((0,0), line, font=font)
            d.text((x+bw/2-(bb[2]-bb[0])/2, yy), line, fill='black', font=font)
            yy += 48
        if i < len(boxes)-1:
            ax = x + bw + 7
            ay = y0 + bh/2
            d.line([ax, ay, ax+gap-18, ay], fill='black', width=5)
            d.polygon([(ax+gap-18, ay-13), (ax+gap-18, ay+13), (ax+gap-2, ay)], fill='black')
    img.save(path)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run('– ' + text)
    set_run_font(r, 14)
    return p


def page_break(doc):
    doc.add_page_break()


def build():
    interface = ASSETS / 'ascon_interface.png'
    drawing = ASSETS / 'shaft_drawing.png'
    flow = ASSETS / 'flowchart.png'
    download(INTERFACE_URL, interface)
    download(DRAWING_URL, drawing)
    create_flowchart(flow)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(14)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5

    # Title page
    add_center(doc, 'Образовательная автономная некоммерческая организация', size=14)
    add_center(doc, 'высшего образования', size=14)
    add_center(doc, '«МОСКОВСКИЙ ТЕХНОЛОГИЧЕСКИЙ ИНСТИТУТ»', size=14, bold=True)
    for _ in range(8):
        doc.add_paragraph()
    add_center(doc, 'ОТЧЕТ', bold=True, size=14)
    add_center(doc, 'о прохождении учебной практики', bold=True, size=14)
    add_center(doc, 'по профессиональному модулю ПМ.01', bold=True, size=14)
    add_center(doc, '«Ведение процесса чертежных и простых', bold=True, size=14)
    add_center(doc, 'расчетно-конструкторских работ»', bold=True, size=14)
    doc.add_paragraph()
    add_center(doc, 'обучающегося группы ОЗПЧКо-25091р', size=14)
    add_center(doc, 'Трубачевой Екатерины Вадимовны', size=14)
    for _ in range(9):
        doc.add_paragraph()
    add_right(doc, 'Дата: ____________________', size=14)
    add_right(doc, '_____________________________', size=14)
    add_right(doc, '(подпись обучающегося)', size=14)

    page_break(doc)
    add_center(doc, 'СОДЕРЖАНИЕ', bold=True)
    doc.add_paragraph()
    toc = [
        ('ВВЕДЕНИЕ', '3'),
        ('1. ВИДЫ ПРОГРАММНОГО ОБЕСПЕЧЕНИЯ ДЛЯ СОЗДАНИЯ ЧЕРТЕЖЕЙ И ВЫПОЛНЕНИЯ РАСЧЕТОВ', '4'),
        ('1.1. Назначение и классификация САПР', '4'),
        ('1.2. Сравнительный анализ программных решений', '5'),
        ('2. ВОЗМОЖНОСТИ СИСТЕМЫ ТРЕХМЕРНОГО ПРОЕКТИРОВАНИЯ КОМПАС-3D', '6'),
        ('2.1. Назначение системы и типы документов', '6'),
        ('2.2. Основные элементы интерфейса', '7'),
        ('2.3. Инструменты моделирования и оформления чертежей', '8'),
        ('2.4. Библиотеки и прикладные расширения', '9'),
        ('2.5. Пример последовательности выполнения прикладной задачи', '10'),
        ('ЗАКЛЮЧЕНИЕ', '11'),
        ('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', '12'),
        ('ПРИЛОЖЕНИЕ А. Пример рабочего чертежа детали', '13'),
    ]
    for title, num in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Cm(16), alignment=2, leader=1)
        r = p.add_run(title)
        set_run_font(r, 14)
        r2 = p.add_run('\t' + num)
        set_run_font(r2, 14)

    new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    new_sec.top_margin = Cm(2)
    new_sec.bottom_margin = Cm(2)
    new_sec.left_margin = Cm(3)
    new_sec.right_margin = Cm(1)
    add_page_number(new_sec, 3)

    # Page 3
    add_heading(doc, 'ВВЕДЕНИЕ')
    add_body(doc, 'Современная конструкторская деятельность связана с подготовкой графической документации, построением моделей изделий и проверкой принятых технических решений. При ручном выполнении таких операций значительная часть времени затрачивается на оформление, внесение повторяющихся изменений и согласование взаимосвязанных видов. Системы автоматизированного проектирования позволяют организовать эти процессы последовательно и снизить вероятность технических ошибок.')
    add_body(doc, 'Цель учебной практики — рассмотреть возможности системы трехмерного проектирования КОМПАС-3D применительно к выполнению чертежных и простых расчетно-конструкторских работ. Для достижения цели необходимо охарактеризовать основные классы САПР, сопоставить распространенные программные решения, изучить интерфейс КОМПАС-3D и показать последовательность подготовки модели и рабочего чертежа детали.')
    add_body(doc, 'КОМПАС-3D разработан компанией АСКОН и применяется для создания трехмерных моделей, чертежей, спецификаций и связанных с ними конструкторских документов. Особое значение имеет поддержка ассоциативной связи между моделью и чертежом: при корректировке геометрии изделия можно актуализировать производные изображения и снизить объем повторного редактирования документации.')
    add_body(doc, 'В отчете рассмотрены функции, которые наиболее непосредственно связаны с задачами учебной практики: создание параметрических эскизов, формирование объемной геометрии, получение ассоциативных видов, нанесение размеров, использование библиотек стандартных элементов и подготовка итогового чертежа.')

    # Page 4
    page_break(doc)
    add_heading(doc, '1. ВИДЫ ПРОГРАММНОГО ОБЕСПЕЧЕНИЯ ДЛЯ СОЗДАНИЯ ЧЕРТЕЖЕЙ И ВЫПОЛНЕНИЯ РАСЧЕТОВ')
    add_heading(doc, '1.1. Назначение и классификация САПР')
    add_body(doc, 'Система автоматизированного проектирования представляет собой совокупность программных, информационных и технических средств, предназначенных для поддержки проектных работ. В инженерной практике САПР используются для построения геометрических моделей, выпуска документации, выполнения расчетов и организации данных об изделии.')
    add_body(doc, 'По характеру решаемых задач можно выделить несколько групп программных решений:')
    add_bullet(doc, 'системы двумерного проектирования, ориентированные на создание и редактирование чертежей;')
    add_bullet(doc, 'системы трехмерного твердотельного и поверхностного моделирования;')
    add_bullet(doc, 'специализированные решения для архитектурного, строительного и технологического проектирования;')
    add_bullet(doc, 'расчетные комплексы, предназначенные для проверки прочности, устойчивости, тепловых и иных характеристик;')
    add_bullet(doc, 'средства управления инженерными данными, обеспечивающие хранение версий и согласование документов.')
    add_body(doc, 'В практической деятельности указанные направления нередко объединяются в единой программной среде. Такой подход позволяет использовать одну геометрическую основу при подготовке различных документов и уменьшает количество несогласованных изменений.')

    # Page 5
    page_break(doc)
    add_heading(doc, '1.2. Сравнительный анализ программных решений')
    add_body(doc, 'Для выбора инструмента необходимо учитывать назначение проекта, требования к оформлению документации, доступные библиотеки, совместимость форматов и возможность дальнейшей передачи модели между участниками процесса. В таблице 1 приведено сопоставление нескольких распространенных систем.')
    add_right(doc, 'Таблица 1', size=14)
    add_center(doc, 'Сравнительный анализ программного обеспечения для проектных работ', size=14)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], 'Программное решение', True)
    set_cell_text(hdr[1], 'Основное назначение', True)
    set_cell_text(hdr[2], 'Характерные возможности', True)
    rows = [
        ('КОМПАС-3D', 'Инженерное 2D- и 3D-проектирование', 'Модели деталей и сборок, ассоциативные чертежи, спецификации, библиотеки стандартных изделий.'),
        ('nanoCAD Механика PRO', 'Машиностроительное проектирование', 'Инструменты 2D-проектирования и 3D-моделирования, оформление документации по отраслевым стандартам.'),
        ('AutoCAD', 'Универсальное проектирование и черчение', 'Создание и редактирование 2D-чертежей, работа с 3D-геометрией, обмен данными.'),
        ('SolidWorks', 'Трехмерное проектирование изделий', 'Параметрическое моделирование, сборки, выпуск документации и средства инженерного анализа.'),
        ('Tekla Structures', 'Информационное моделирование строительных конструкций', 'Детализированные модели конструкций, выпуск документации, координация проектных данных.'),
    ]
    for item in rows:
        cells = table.add_row().cells
        for i, text in enumerate(item):
            set_cell_text(cells[i], text)
        prevent_row_split(table.rows[-1])
    set_repeat_table_header(table.rows[0])
    set_col_widths(table, [3.4, 4.2, 9.0])
    add_body(doc, 'КОМПАС-3D целесообразно рассматривать как универсальный инструмент для выполнения учебных задач, связанных с моделированием деталей и оформлением чертежей. Система позволяет последовательно пройти путь от эскиза до графического документа, сохраняя связь между этапами работы.')

    # Page 6
    page_break(doc)
    add_heading(doc, '2. ВОЗМОЖНОСТИ СИСТЕМЫ ТРЕХМЕРНОГО ПРОЕКТИРОВАНИЯ КОМПАС-3D')
    add_heading(doc, '2.1. Назначение системы и типы документов')
    add_body(doc, 'КОМПАС-3D предназначен для автоматизации проектно-конструкторских работ. В рамках одной среды пользователь может создать модель детали, собрать изделие из компонентов, подготовить чертеж и оформить сопутствующие документы. Такой подход особенно удобен при внесении изменений, поскольку модель выступает исходной информационной основой проекта.')
    add_body(doc, 'Основные типы документов, применяемые при выполнении простых расчетно-конструкторских работ, представлены в таблице 2.')
    add_right(doc, 'Таблица 2', size=14)
    add_center(doc, 'Основные типы документов КОМПАС-3D', size=14)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    set_cell_text(table.rows[0].cells[0], 'Тип документа', True)
    set_cell_text(table.rows[0].cells[1], 'Назначение', True)
    for a,b in [
        ('Деталь', 'Создание трехмерной модели отдельного изделия.'),
        ('Сборка', 'Объединение компонентов и проверка взаимного расположения деталей.'),
        ('Чертеж', 'Оформление графического документа с видами, разрезами, размерами и обозначениями.'),
        ('Фрагмент', 'Подготовка вспомогательных графических построений без основной надписи.'),
        ('Спецификация', 'Формирование перечня составных частей изделия.'),
    ]:
        cells = table.add_row().cells
        set_cell_text(cells[0], a)
        set_cell_text(cells[1], b)
        prevent_row_split(table.rows[-1])
    set_repeat_table_header(table.rows[0])
    set_col_widths(table, [4.2, 12.4])
    add_body(doc, 'При выполнении учебной задачи ключевыми являются документы «Деталь» и «Чертеж». Сначала формируется объемная геометрия изделия, затем на ее основе создаются ассоциативные виды и выполняется оформление графического документа.')

    # Page 7
    page_break(doc)
    add_heading(doc, '2.2. Основные элементы интерфейса')
    add_body(doc, 'После создания или открытия документа в окне КОМПАС-3D отображаются элементы, предназначенные для выбора команд, управления объектами и работы с геометрией. В официальной справочной системе АСКОН выделены инструментальная область, строка вкладок, главное меню, поле поиска команд, панели управления и графическая область.')
    add_body(doc, 'Графическая область занимает основную часть окна и используется для отображения модели или чертежа. Инструментальная область обеспечивает доступ к командам создания и редактирования объектов. Панели управления позволяют работать с деревом документа, параметрами операций и другими связанными данными. Общая структура интерфейса показана на рисунке 1.')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    r.add_picture(str(interface), width=Cm(15.8))
    add_caption(doc, 'Рисунок 1 – Основные элементы интерфейса КОМПАС-3D')
    add_body(doc, 'Наличие поля поиска команд позволяет быстро найти нужную операцию по ее названию. Это удобно при освоении системы: пользователь может последовательно изучать назначение инструментов без необходимости запоминать расположение каждой команды в меню.')

    # Page 8
    page_break(doc)
    add_heading(doc, '2.3. Инструменты моделирования и оформления чертежей')
    add_body(doc, 'Создание модели детали обычно начинается с построения эскиза. В эскизе задаются основные геометрические элементы: отрезки, окружности, дуги, осевые линии и зависимости между объектами. Размеры определяют положение и форму элементов, а параметрические связи позволяют сохранить требуемую геометрию при изменении отдельных значений.')
    add_body(doc, 'После завершения эскиза применяются операции формообразования. Для деталей типа тел вращения используется операция вращения профиля относительно оси. Для призматических элементов применяется выдавливание. Дополнительные конструктивные особенности могут быть сформированы операциями вырезания, построения отверстий, фасок, скруглений и массивов.')
    add_body(doc, 'На основе готовой модели формируется чертеж. Пользователь размещает главный вид и необходимые проекции, при необходимости добавляет разрезы и сечения. Затем выполняется нанесение размеров, осевых линий, обозначений шероховатости и иных элементов оформления.')
    add_right(doc, 'Таблица 3', size=14)
    add_center(doc, 'Связь этапов моделирования и оформления документации', size=14)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,t in enumerate(['Этап', 'Используемые инструменты', 'Результат']):
        set_cell_text(table.rows[0].cells[i], t, True)
    for row in [
        ('Построение эскиза', 'Геометрические примитивы, размеры, зависимости', 'Параметрический профиль детали'),
        ('Создание модели', 'Вращение, выдавливание, вырезание, отверстия', 'Трехмерная геометрия изделия'),
        ('Подготовка чертежа', 'Ассоциативные виды, разрезы, сечения', 'Набор изображений детали'),
        ('Оформление', 'Размеры, обозначения, технические требования', 'Рабочий чертеж'),
    ]:
        cells = table.add_row().cells
        for i,t in enumerate(row): set_cell_text(cells[i], t)
        prevent_row_split(table.rows[-1])
    set_repeat_table_header(table.rows[0])
    set_col_widths(table, [3.7, 6.2, 6.7])

    # Page 9
    page_break(doc)
    add_heading(doc, '2.4. Библиотеки и прикладные расширения')
    add_body(doc, 'Стандартные возможности КОМПАС-3D могут быть дополнены библиотеками и приложениями. Их использование сокращает количество повторяющихся операций и позволяет обращаться к готовым элементам, которые часто встречаются в инженерной документации. Особенно это важно при проектировании сборок и оформлении чертежей с типовыми компонентами.')
    add_body(doc, 'В зависимости от состава установленного программного обеспечения пользователь может работать с библиотеками стандартных изделий, элементами крепежа, отверстиями и другими объектами. Отдельные приложения предназначены для решения специализированных задач и выполнения инженерных расчетов.')
    add_right(doc, 'Таблица 4', size=14)
    add_center(doc, 'Примеры применения библиотек и расширений', size=14)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(table.rows[0].cells[0], 'Инструмент', True)
    set_cell_text(table.rows[0].cells[1], 'Практическое назначение', True)
    for a,b in [
        ('Библиотеки стандартных изделий', 'Добавление типовых компонентов и сокращение времени на повторное построение.'),
        ('Библиотеки отверстий и конструктивных элементов', 'Быстрое формирование распространенных элементов детали.'),
        ('Средства работы со спецификациями', 'Подготовка перечня составных частей изделия.'),
        ('Расчетные приложения', 'Проверка отдельных характеристик модели в пределах доступного функционала.'),
    ]:
        cells = table.add_row().cells
        set_cell_text(cells[0], a)
        set_cell_text(cells[1], b)
        prevent_row_split(table.rows[-1])
    set_repeat_table_header(table.rows[0])
    set_col_widths(table, [6.0, 10.6])
    add_body(doc, 'Применение библиотек не исключает необходимость проверки результата. Готовый элемент должен соответствовать назначению детали, выбранным размерам и требованиям к оформлению документации.')

    # Page 10
    page_break(doc)
    add_heading(doc, '2.5. Пример последовательности выполнения прикладной задачи')
    add_body(doc, 'В качестве прикладного примера можно рассмотреть подготовку модели и рабочего чертежа ступенчатого вала. Деталь относится к телам вращения, поэтому ее удобно строить на основе продольного профиля. На эскизе задаются участки различного диаметра, длины ступеней и ось вращения. После применения операции вращения формируется объемная модель.')
    add_body(doc, 'На следующем этапе создается документ «Чертеж», в который добавляются ассоциативные виды модели. Для отображения формы отдельных элементов при необходимости используются сечения и выносные изображения. Затем наносятся размеры, обозначения шероховатости и технические требования. Общая последовательность представлена на рисунке 2.')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(flow), width=Cm(16.0))
    add_caption(doc, 'Рисунок 2 – Последовательность подготовки модели и рабочего чертежа')
    add_body(doc, 'Преимущество такой последовательности заключается в логической связи между этапами. Если исходная геометрия корректируется, модель остается основой для актуализации производных изображений. Это снижает риск расхождений между представлением детали и ее чертежом.')

    # Page 11
    page_break(doc)
    add_heading(doc, 'ЗАКЛЮЧЕНИЕ')
    add_body(doc, 'В ходе учебной практики были рассмотрены основные виды программного обеспечения, применяемого для создания чертежей и выполнения расчетно-конструкторских работ. Системы автоматизированного проектирования позволяют объединить построение геометрии, оформление документации и работу с инженерными данными в последовательный процесс.')
    add_body(doc, 'КОМПАС-3D предоставляет инструменты для создания параметрических эскизов, трехмерных моделей деталей и сборок, ассоциативных чертежей и спецификаций. Для решения прикладных задач важны операции вращения и выдавливания, средства построения отверстий, фасок и скруглений, а также инструменты нанесения размеров и обозначений.')
    add_body(doc, 'Интерфейс системы организован таким образом, чтобы пользователь мог работать с документом, геометрией и параметрами операций в единой среде. Панели управления, инструментальная область и поле поиска команд упрощают выбор инструментов и позволяют последовательно осваивать возможности программы.')
    add_body(doc, 'На примере ступенчатого вала показано, что подготовка рабочего чертежа может быть организована как цепочка взаимосвязанных действий: анализ исходных данных, построение эскиза, создание модели, формирование ассоциативных видов и итоговое оформление. Такой подход способствует закреплению навыков чтения чертежей, пониманию логики моделирования и аккуратной подготовке конструкторской документации.')
    add_body(doc, 'По результатам практики закреплены представления о назначении САПР и приобретены навыки анализа инструментов, необходимых для выполнения простых чертежных и расчетно-конструкторских работ.')

    # Page 12
    page_break(doc)
    add_heading(doc, 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ')
    sources = [
        '1. КОМПАС-3D. Официальная страница продукта. — URL: https://kompas.ru/kompas-3d/ (дата обращения: 01.06.2026).',
        '2. КОМПАС-3D v24. Интерфейс системы : официальная справочная система АСКОН. — URL: https://help.ascon.ru/KOMPAS/24/ru-RU/ae1920959.html (дата обращения: 01.06.2026).',
        '3. КОМПАС-3D. Окно системы : официальная справочная система АСКОН. — URL: https://help.ascon.ru/KOMPAS/21/ru-RU/63_1_2_1_okno_sistemy.html (дата обращения: 01.06.2026).',
        '4. КОМПАС-3D v24. Панели управления : официальная справочная система АСКОН. — URL: https://help.ascon.ru/KOMPAS/24/ru-RU/ae1730271.html (дата обращения: 01.06.2026).',
        '5. КОМПАС-3D v24. Панель параметров при выполнении команды : официальная справочная система АСКОН. — URL: https://help.ascon.ru/KOMPAS/24/ru-RU/33_1_1_panelq_svojstv.html (дата обращения: 01.06.2026).',
        '6. nanoCAD Механика PRO. Официальная страница программного продукта. — URL: https://www.nanocad.ru/products/nanocad-mekhanika-pro/ (дата обращения: 01.06.2026).',
        '7. AutoCAD. Официальная страница программного продукта Autodesk. — URL: https://www.autodesk.com/products/autocad/overview (дата обращения: 01.06.2026).',
        '8. SOLIDWORKS Design. Официальная страница программного продукта. — URL: https://www.solidworks.com/product/solidworks-design (дата обращения: 01.06.2026).',
        '9. Проектирование технологического процесса изготовления детали «Вал ступенчатый» // Инженерный портал «В Масштабе.ру». — URL: https://vmasshtabe.ru/mashinostroenie-i-mehanika/tm/proektirovanie-tehnologicheskogo-protsessa-izgotovleniya-detali-val-stupenchatyiy.html (дата обращения: 01.06.2026).',
    ]
    for s in sources:
        add_body(doc, s)

    # Page 13
    page_break(doc)
    add_heading(doc, 'ПРИЛОЖЕНИЕ А')
    add_center(doc, 'Пример рабочего чертежа детали «Вал ступенчатый»', bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(drawing), width=Cm(16.2))
    add_caption(doc, 'Рисунок А.1 – Рабочий чертеж детали «Вал ступенчатый»')
    add_body(doc, 'Источник: инженерный портал «В Масштабе.ру». Чертеж приведен в качестве примера оформления графической документации и иллюстрации результата, который может быть подготовлен средствами САПР.')

    # Remove accidental empty footer from first section
    for p in doc.sections[0].footer.paragraphs:
        p.text = ''

    doc.save(DOCX)
    with zipfile.ZipFile(DOCX, 'r') as z:
        bad = z.testzip()
        if bad:
            raise RuntimeError(f'DOCX archive error: {bad}')
    print(DOCX)

if __name__ == '__main__':
    build()
